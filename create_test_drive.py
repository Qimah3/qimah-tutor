"""
Generate synthetic 'test drive' folder for integration tests.
Produces enough content for >30 chunks (chunk_size=500, overlap=50).
"""
import os
import sys

# Add venv to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "venv/lib/python3.12/site-packages"))

import fitz  # PyMuPDF
import docx
from PIL import Image, ImageDraw, ImageFont

DEST = os.path.join(os.path.dirname(__file__), "test drive")
os.makedirs(DEST, exist_ok=True)

# ── PDF: 10 pages × ~1,500 chars of CS101-style text → ~40 chunks ──────────
CS101_PAGE = """\
CS101 - Introduction to Computer Science
Chapter {n}: Data Structures and Algorithms

Arrays and ArrayLists store elements in contiguous memory. The indexOf() method
searches linearly and returns the first matching index, or -1 if not found.
Time complexity is O(n) in the worst case.

Sorting algorithms: Bubble Sort compares adjacent elements and swaps them if
out of order. Selection Sort finds the minimum element and places it at the
beginning. Both have O(n^2) time complexity in the worst and average cases.

Binary Search requires a sorted array. It compares the target to the middle
element and discards half the array on each step. Time complexity: O(log n).
This is significantly faster than linear search for large datasets.

Linked Lists: Each node holds a value and a pointer to the next node. Insertion
at the head is O(1). Searching requires O(n) traversal. Doubly linked lists add
a previous pointer, allowing backward traversal and easier deletion.

Stacks (LIFO) and Queues (FIFO) are abstract data types. A stack supports push
and pop. A queue supports enqueue and dequeue. Both can be implemented with
arrays or linked lists. Stacks are used in recursion and expression evaluation.

Recursion: A function that calls itself with a smaller subproblem. Base case
stops the recursion. Examples include factorial, Fibonacci, and tree traversal.
Each recursive call adds a frame to the call stack. Deep recursion risks overflow.
"""  # ~1,400 chars per page

pdf_path = os.path.join(DEST, "cs101_lecture_notes.pdf")
doc = fitz.open()
for i in range(1, 11):  # 10 pages
    page = doc.new_page(width=595, height=842)
    page.insert_text(
        (50, 60),
        CS101_PAGE.format(n=i),
        fontname="helv",
        fontsize=10,
        color=(0, 0, 0),
    )
doc.save(pdf_path)
print(f"Created PDF: {pdf_path} ({doc.page_count} pages)")

# ── DOCX: Long paragraphs → ~8 chunks ───────────────────────────────────────
DOCX_PARAGRAPHS = [
    "Quiz 4 - CS101 Fall 2025\nStudent ID: 215110365\nDate: October 2025",
    "Question 1: What is the time complexity of binary search on a sorted array of n elements? "
    "Answer: O(log n). Binary search divides the search interval in half with each step, "
    "comparing the target to the middle element. If the target matches, return the index. "
    "If the target is less, search the left half. If greater, search the right half. "
    "This halving continues until the target is found or the interval is empty.",
    "Question 2: Explain the difference between a stack and a queue. "
    "Answer: A stack is a Last In First Out (LIFO) data structure. The last element pushed "
    "is the first element popped. Common operations are push, pop, and peek. "
    "A queue is a First In First Out (FIFO) structure. The first element enqueued is the "
    "first element dequeued. Common operations are enqueue, dequeue, and front. "
    "Stacks are used in function call management, undo operations, and expression parsing. "
    "Queues are used in breadth-first search, task scheduling, and print spooling.",
    "Question 3: Write a Java method that uses indexOf to find the position of a character. "
    "public static int findChar(String s, char c) { return s.indexOf(c); } "
    "The String.indexOf(char ch) method returns the index of the first occurrence of ch, "
    "or -1 if the character is not found in the string. It performs linear search internally.",
    "Question 4: What are the advantages of a doubly linked list over a singly linked list? "
    "Answer: A doubly linked list maintains both next and previous pointers at each node. "
    "This allows backward traversal without restarting from the head. Deletion of a node "
    "given only the node reference is O(1) in a doubly linked list (since we have the previous "
    "pointer), but O(n) in a singly linked list because we must find the predecessor. "
    "The tradeoff is extra memory per node for the second pointer.",
    "Question 5: Describe recursion and identify the base case for computing factorial. "
    "Answer: Recursion is a programming technique where a function solves a problem by "
    "calling itself with a smaller version of the same problem. The base case is the "
    "condition that stops the recursion. For factorial: base case is factorial(0) = 1 "
    "and factorial(1) = 1. Recursive case: factorial(n) = n * factorial(n-1). "
    "Without a base case, recursion would continue infinitely causing a stack overflow.",
    "Bonus: Analyze the sortedArray insertion sort implementation step by step for the "
    "input array [5, 3, 8, 1, 9, 2]. Step 1: [3, 5, 8, 1, 9, 2]. Step 2: [3, 5, 8, 1, 9, 2]. "
    "Step 3: [1, 3, 5, 8, 9, 2]. Step 4: [1, 3, 5, 8, 9, 2]. Step 5: [1, 2, 3, 5, 8, 9]. "
    "The final sorted array is [1, 2, 3, 5, 8, 9]. Insertion sort has O(n^2) worst case "
    "but O(n) best case when the input is already sorted.",
]

docx_path = os.path.join(DEST, "Quiz4_215110365.docx")
document = docx.Document()
document.add_heading("CS101 Quiz 4 Answers", 0)
for para in DOCX_PARAGRAPHS:
    document.add_paragraph(para)
document.save(docx_path)
print(f"Created DOCX: {docx_path}")

# ── Images: Simple text on white background ──────────────────────────────────
IMAGE_TEXTS = [
    ("lab6.jpg", "Lab 6 - CS101\nArrayList and LinkedList Operations\n"
     "Exercise 1: Implement indexOf for ArrayList\n"
     "Exercise 2: Implement insert at head for LinkedList\n"
     "Use Java standard library where appropriate.\n"
     "Submit your solution on the course portal.\n"
     "Due date: Week 10, Sunday midnight."),
    ("lecture_slide.png", "CS101 Lecture - Searching Algorithms\n"
     "Linear Search: O(n) - checks each element\n"
     "Binary Search: O(log n) - sorted arrays only\n"
     "Hash Search: O(1) average - uses hash table\n"
     "Choose the right algorithm based on data size."),
]

for filename, text in IMAGE_TEXTS:
    img = Image.new("RGB", (800, 400), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.multiline_text((40, 40), text, fill=(0, 0, 0), spacing=8)
    img.save(os.path.join(DEST, filename))
    print(f"Created image: {filename}")

print("\nDone! Files in 'test drive/':")
for f in sorted(os.listdir(DEST)):
    size = os.path.getsize(os.path.join(DEST, f))
    print(f"  {f} ({size:,} bytes)")

# Estimate chunk count
# PDF: 10 pages × ~1400 chars → step=450 → ~4 chunks/page → 40 chunks
# DOCX: ~4000 chars total → ~9 chunks
# Images: OCR may or may not extract text
print("\nEstimated chunks from PDF+DOCX alone: ~49 (well above 30)")
